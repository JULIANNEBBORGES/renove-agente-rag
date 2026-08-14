"""
Ingestão e chunking dos documentos da Renove Lavanderias.

Agente focado no colaborador interno (não trata de atendimento a cliente externo).
Fontes de conhecimento textual: regras de negócio (assiduidade/premiação), BPMN do
processo de extração mensal, tutorial de extração do Renov.net e guias técnicos de
treinamento (docs_treinamento/).

Lê esses documentos, extrai o conteúdo textual e organiza em "chunks" (pedaços) com
metadados, prontos para serem transformados em embeddings (Google Gemini) e indexados
num vector store local (FAISS ou ChromaDB).

A planilha de desempenho (Av_Desempenho_Renove.xlsx) NÃO passa por este script —
ela é carregada separadamente com pandas dentro do próprio app Streamlit,
pois serve para responder perguntas numéricas, não perguntas de texto.

Uso:
    python ingest.py

Saída:
    chunks.jsonl — um chunk por linha, formato:
    {"id": "...", "source": "...", "categoria": "...", "texto": "..."}
"""

import json
import re
from pathlib import Path

from docx import Document as DocxDocument
from bs4 import BeautifulSoup

BASE_DIR = Path(__file__).resolve().parent.parent
DOCUMENTOS_DIR = BASE_DIR / "documentos"
TREINAMENTO_DIR = BASE_DIR / "docs_treinamento"
OUTPUT_FILE = Path(__file__).resolve().parent / "chunks.jsonl"

CHUNK_MAX_PALAVRAS = 220
CHUNK_OVERLAP_PALAVRAS = 30


def extrair_texto_docx(caminho: Path) -> str:
    """Extrai parágrafos e tabelas de um arquivo .docx, preservando quebras lógicas."""
    doc = DocxDocument(str(caminho))
    partes = []

    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto:
            # Marca títulos com # pra ajudar o chunking a respeitar seções
            estilo = para.style.name if para.style is not None else ""
            if estilo.startswith("Heading") or estilo == "Title":
                partes.append(f"\n## {texto}\n")
            else:
                partes.append(texto)

    for tabela in doc.tables:
        linhas = []
        for row in tabela.rows:
            celulas = [c.text.strip() for c in row.cells]
            linhas.append(" | ".join(celulas))
        if linhas:
            partes.append("\n".join(linhas))

    return "\n".join(partes)


def extrair_texto_html(caminho: Path) -> str:
    """Extrai texto visível de um HTML, incluindo rótulos de texto dentro de SVG (diagramas BPMN)."""
    with open(caminho, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    partes = []
    for tag in soup.find_all(["h1", "h2", "p"]):
        texto = tag.get_text(strip=True)
        if texto:
            if tag.name in ("h1", "h2"):
                partes.append(f"\n## {texto}\n")
            else:
                partes.append(texto)

    # Rótulos de texto dentro do SVG (nomes de tarefas, lojas, sistema, etc.)
    svg_textos = [t.get_text(strip=True) for t in soup.find_all("text")]
    svg_textos = [t for t in svg_textos if t]
    if svg_textos:
        partes.append("Elementos do diagrama: " + " | ".join(svg_textos))

    return "\n".join(partes)


def extrair_texto_md(caminho: Path) -> str:
    with open(caminho, "r", encoding="utf-8") as f:
        return f.read()


def dividir_em_chunks(texto: str, max_palavras=CHUNK_MAX_PALAVRAS, overlap=CHUNK_OVERLAP_PALAVRAS):
    """Chunking simples por parágrafo/seção, respeitando um teto de palavras com overlap."""
    # Quebra em blocos por linha em branco dupla ou cabeçalho de seção
    blocos = re.split(r"\n(?=## )|\n\s*\n", texto)
    blocos = [b.strip() for b in blocos if b.strip()]

    chunks = []
    atual = []
    contagem = 0

    for bloco in blocos:
        n_palavras = len(bloco.split())
        if contagem + n_palavras > max_palavras and atual:
            chunks.append("\n\n".join(atual))
            # overlap: mantém o último bloco pro próximo chunk, se couber
            if overlap > 0 and atual:
                ultimo = atual[-1]
                atual = [ultimo] if len(ultimo.split()) <= overlap else []
                contagem = len(atual[0].split()) if atual else 0
            else:
                atual = []
                contagem = 0
        atual.append(bloco)
        contagem += n_palavras

    if atual:
        chunks.append("\n\n".join(atual))

    return chunks


def processar_arquivo(caminho: Path, categoria: str):
    sufixo = caminho.suffix.lower()
    if sufixo == ".docx":
        texto = extrair_texto_docx(caminho)
    elif sufixo in (".html", ".htm"):
        texto = extrair_texto_html(caminho)
    elif sufixo == ".md":
        texto = extrair_texto_md(caminho)
    else:
        print(f"Aviso: tipo de arquivo não suportado, ignorando -> {caminho.name}")
        return []

    chunks_texto = dividir_em_chunks(texto)
    registros = []
    for i, chunk in enumerate(chunks_texto):
        registros.append({
            "id": f"{caminho.stem}_{i:03d}",
            "source": caminho.name,
            "categoria": categoria,
            "texto": chunk,
        })
    return registros


def main():
    todos_chunks = []

    # Agente focado no colaborador interno — atendimento_ao_cliente está fora de escopo.
    fontes = [
        (DOCUMENTOS_DIR / "gestao_processos_internos", "gestao_processos_internos"),
        (TREINAMENTO_DIR, "docs_treinamento"),
    ]

    for pasta, categoria in fontes:
        if not pasta.exists():
            continue
        for arquivo in sorted(pasta.iterdir()):
            if arquivo.is_file() and arquivo.suffix.lower() in (".docx", ".html", ".htm", ".md"):
                registros = processar_arquivo(arquivo, categoria)
                todos_chunks.extend(registros)
                print(f"{arquivo.name}: {len(registros)} chunks")

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        for chunk in todos_chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    print(f"\nTotal: {len(todos_chunks)} chunks salvos em {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
